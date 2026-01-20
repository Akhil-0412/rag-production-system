import os
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from pathlib import Path

# File handling libraries
import pypdf
import docx

@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

class FileLoader:
    """
    Handles loading of various file formats: PDF, DOCX, TXT, MD.
    """
    
    def load_file(self, file_path: str) -> Optional[Document]:
        """
        Detects file extension and routes to appropriate loader.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        extension = path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._load_pdf(file_path)
            elif extension == '.docx':
                return self._load_docx(file_path)
            elif extension in ['.txt', '.md', '.markdown']:
                return self._load_text(file_path)
            else:
                print(f"Unsupported file type: {extension}")
                return None
        except Exception as e:
            print(f"Error loading file {file_path}: {e}")
            return None

    def _load_pdf(self, file_path: str) -> Document:
        text_content = []
        metadata = {
            "source": os.path.basename(file_path),
            "file_path": file_path,
            "type": "pdf",
            "page_count": 0
        }
        
        with open(file_path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            metadata["page_count"] = len(reader.pages)
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    # We could store page numbers in chunks later, 
                    # but for raw loading we treat it as a stream or join with markers.
                    # For now, let's join with a pagebreak marker to allow chunking keying.
                    text_content.append(text)
        
        full_text = "\n\n".join(text_content)
        return Document(content=full_text, metadata=metadata)

    def _load_docx(self, file_path: str) -> Document:
        doc = docx.Document(file_path)
        text_content = [para.text for para in doc.paragraphs if para.text.strip()]
        
        metadata = {
            "source": os.path.basename(file_path),
            "file_path": file_path,
            "type": "docx"
        }
        
        return Document(content="\n\n".join(text_content), metadata=metadata)

    def _load_text(self, file_path: str) -> Document:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            
        metadata = {
            "source": os.path.basename(file_path),
            "file_path": file_path,
            "type": "text"
        }
        return Document(content=content, metadata=metadata)

    def load_directory(self, directory_path: str) -> List[Document]:
        documents = []
        path = Path(directory_path)
        
        if not path.exists():
            print(f"Directory not found: {directory_path}")
            return []

        for file_path in path.glob("**/*"):
            if file_path.is_file():
                doc = self.load_file(str(file_path))
                if doc:
                    documents.append(doc)
                    
        return documents
